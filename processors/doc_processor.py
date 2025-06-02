from docx import Document
from typing import Dict, Any
from .base_processor import BaseProcessor

class DocProcessor(BaseProcessor):
    """Process DOC and DOCX files"""
    
    def process(self, filepath: str) -> Dict[str, Any]:
        """Extract text and structure from Word documents"""
        result = {
            'text': '',
            'paragraphs': [],
            'tables': [],
            'metadata': {}
        }
        
        try:
            doc = Document(filepath)
            
            # Extract metadata
            core_props = doc.core_properties
            result['metadata'] = {
                'title': core_props.title or '',
                'author': core_props.author or '',
                'subject': core_props.subject or '',
                'created': str(core_props.created) if core_props.created else '',
                'modified': str(core_props.modified) if core_props.modified else ''
            }
            
            full_text = []
            
            # Extract paragraphs
            for para in doc.paragraphs:
                if para.text.strip():
                    para_data = {
                        'text': para.text.strip(),
                        'style': para.style.name if para.style else 'Normal'
                    }
                    result['paragraphs'].append(para_data)
                    
                    # Format based on style
                    if para.style and 'Heading' in para.style.name:
                        level = self._get_heading_level(para.style.name)
                        full_text.append(f"{'#' * level} {para.text.strip()}")
                    else:
                        full_text.append(para.text.strip())
            
            # Extract tables
            for table_idx, table in enumerate(doc.tables):
                table_data = {
                    'table_index': table_idx,
                    'rows': []
                }
                
                for row in table.rows:
                    row_data = []
                    for cell in row.cells:
                        row_data.append(cell.text.strip())
                    table_data['rows'].append(row_data)
                
                result['tables'].append(table_data)
                
                # Add table as markdown to text
                table_md = self._table_to_markdown(table_data['rows'])
                full_text.append(f"\n--- Table {table_idx + 1} ---\n{table_md}")
            
            result['text'] = '\n\n'.join(full_text)
            
        except Exception as e:
            result['error'] = str(e)
            result['text'] = f"Error processing Word document: {str(e)}"
        
        return result
    
    def _get_heading_level(self, style_name: str) -> int:
        """Extract heading level from style name"""
        if 'Heading 1' in style_name:
            return 1
        elif 'Heading 2' in style_name:
            return 2
        elif 'Heading 3' in style_name:
            return 3
        elif 'Heading 4' in style_name:
            return 4
        elif 'Heading 5' in style_name:
            return 5
        elif 'Heading 6' in style_name:
            return 6
        return 2  # Default
    
    def _table_to_markdown(self, rows):
        """Convert table rows to markdown format"""
        if not rows or not rows[0]:
            return ""
        
        markdown = ""
        
        # Header row
        headers = rows[0]
        markdown += "| " + " | ".join(headers) + " |\n"
        markdown += "| " + " | ".join(["---"] * len(headers)) + " |\n"
        
        # Data rows
        for row in rows[1:]:
            # Pad row to match header length
            while len(row) < len(headers):
                row.append("")
            markdown += "| " + " | ".join(row) + " |\n"
        
        return markdown
    
    def to_markdown(self, content: Dict[str, Any]) -> str:
        """Convert Word document content to markdown"""
        if isinstance(content, str):
            return content
        
        markdown = ""
        
        # Add metadata
        if content.get('metadata'):
            metadata = content['metadata']
            markdown += "# Document Information\n\n"
            for key, value in metadata.items():
                if value:
                    markdown += f"**{key.title()}:** {value}\n\n"
        
        # Add main text content
        if content.get('text'):
            markdown += "# Content\n\n"
            markdown += content['text']
        
        return markdown
